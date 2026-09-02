#!/usr/bin/env python3
"""JSON AutoQA items → bilingual Chinese / live-language Word review (.docx)."""

from __future__ import annotations

import argparse
import json
import sys
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from common import LANG_SHORT, SITE_LABEL, ensure_in_product_output
from gate import assert_can_deliver, product_dir_from_path

ROOT = Path(__file__).resolve().parents[1]


def _font(run, *, east_asia: str = "宋体", ascii_font: str = "Calibri", size: int = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asia)


def _p(doc, text: str, *, size: int = 11, bold: bool = False, space_after: int = 6, color: RGBColor | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _font(run, size=size, bold=bold, color=color)
    return p


def _set_cell(cell, label: str, body: str, *, label_color: RGBColor) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r0 = p.add_run(label + "\n")
    _font(r0, size=9, bold=True, color=label_color)
    r1 = p.add_run(body)
    _font(r1, size=11)


def load_data(path: Path) -> dict:
    data = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(data, list):
        return {"items": data}
    return data


def ensure_under_output(path: Path) -> Path:
    return ensure_in_product_output(path, ROOT / "output")


def main() -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--input", required=True)
    p.add_argument("--output", required=True)
    p.add_argument(
        "--gate-ok",
        action="store_true",
        help="Internal: allow only after run_autoqa.py deliver validated",
    )
    args = p.parse_args()

    qa_json = Path(args.input).resolve()
    if not args.gate_ok:
        raise SystemExit(
            "Direct write blocked. Use:\n"
            f"  python scripts/run_autoqa.py deliver --qa-json {qa_json}"
        )
    product_dir = product_dir_from_path(qa_json, ROOT / "output")
    assert_can_deliver(product_dir, qa_json)

    data = load_data(qa_json)
    items = data.get("items") or []
    if not items:
        raise SystemExit("no items")
    missing = [
        i
        for i, row in enumerate(items)
        if not str(row.get("question_zh") or "").strip() or not str(row.get("answer_zh") or "").strip()
    ]
    if missing:
        raise SystemExit(f"missing question_zh/answer_zh on items {missing[:8]}{'…' if len(missing) > 8 else ''}")

    out = ensure_under_output(Path(args.output))
    out.parent.mkdir(parents=True, exist_ok=True)

    labels = data.get("variant_labels") or {}
    title = data.get("title") or "AutoQA"
    mp = str(items[0].get("marketplace") or "US").upper()
    site = SITE_LABEL.get(mp, mp)

    grouped: OrderedDict[str, list] = OrderedDict()
    for row in items:
        grouped.setdefault(str(row["asin"]).upper(), []).append(row)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = h.add_run(f"AutoQA 人工核查稿（中英对照）· {title}")
    _font(r, size=16, bold=True)
    h.paragraph_format.space_after = Pt(8)

    _p(doc, f"站点：{site}（{mp}）　合计：{len(items)} 条", size=11, space_after=4)
    _p(
        doc,
        "上栏中文供核对意思与避重就轻；下栏是问答区成稿（英/西等）。改下栏才会进服务商表。中文不上架。",
        size=10,
        color=RGBColor(0x55, 0x55, 0x55),
        space_after=12,
    )

    counts = "　".join(
        f"{labels.get(asin, asin)} {len(rows)}条" for asin, rows in grouped.items()
    )
    _p(doc, counts, size=10, space_after=14)

    n = 0
    mute = RGBColor(0x88, 0x66, 0x22)
    ink = RGBColor(0x33, 0x33, 0x33)
    for asin, rows in grouped.items():
        variant = labels.get(asin, "")
        head = f"{variant} · {asin}" if variant else asin
        _p(doc, head, size=13, bold=True, space_after=8)
        for row in rows:
            n += 1
            lang = str(row.get("language") or "")
            tag = LANG_SHORT.get(lang, lang or "原文")
            note = str(row.get("note") or "").strip()
            meta = f"{n}. {note or '—'}　{tag}"
            _p(doc, meta, size=10, bold=True, space_after=4, color=ink)

            table = doc.add_table(rows=4, cols=1)
            table.style = "Table Grid"
            _set_cell(table.cell(0, 0), "问 · 中", str(row.get("question_zh") or "").strip(), label_color=mute)
            _set_cell(table.cell(1, 0), f"问 · {tag}", str(row.get("question") or "").strip(), label_color=mute)
            _set_cell(table.cell(2, 0), "答 · 中", str(row.get("answer_zh") or "").strip(), label_color=mute)
            _set_cell(table.cell(3, 0), f"答 · {tag}", str(row.get("answer") or "").strip(), label_color=mute)
            doc.add_paragraph().paragraph_format.space_after = Pt(10)

    doc.save(out)
    print(f"wrote {out} ({n} QAs, bilingual review)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
